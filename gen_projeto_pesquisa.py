"""Gera o Projeto de Pesquisa em .docx seguindo o template USP/ESALQ."""
from docx import Document
from docx.shared import Pt, Emu, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Margens (mesmo do template: 900430 EMU ≈ 2.36 cm) ──
for section in doc.sections:
    section.top_margin = 900430
    section.bottom_margin = 900430
    section.left_margin = 900430
    section.right_margin = 900430

# ── Estilo padrão: Arial 11 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# ── Helpers ──
def add_heading_bold(text, alignment=None, spacing_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = spacing_after
    p.paragraph_format.space_before = Pt(6)
    return p

def add_body(text, spacing_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = spacing_after
    return p

def add_body_runs(runs_list, spacing_after=Pt(6)):
    """runs_list: [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    for text, bold, italic in runs_list:
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = spacing_after
    return p

def add_ref(text):
    """Referência: alinhado à esquerda, espaçamento simples, sem negrito."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    return p

# ═══════════════════════════════════════════════════════════
# PRÉ-TEXTUAIS
# ═══════════════════════════════════════════════════════════

p = doc.add_paragraph()
run = p.add_run('Aluno: ')
run.bold = True
run.font.name = 'Arial'
run.font.size = Pt(11)
run = p.add_run('Maycon Henrique Aranha Da Silva')
run.font.name = 'Arial'
run.font.size = Pt(11)
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph()
run = p.add_run('Orientador(a): ')
run.bold = True
run.font.name = 'Arial'
run.font.size = Pt(11)
run = p.add_run('[a definir]')
run.font.name = 'Arial'
run.font.size = Pt(11)
p.paragraph_format.line_spacing = 1.5

p = doc.add_paragraph()
run = p.add_run('Curso: ')
run.bold = True
run.font.name = 'Arial'
run.font.size = Pt(11)
run = p.add_run('MBA em Data Science e Analytics')
run.font.name = 'Arial'
run.font.size = Pt(11)
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(18)

# ── TÍTULO ──
add_heading_bold(
    'Determinantes da insatisfação do consumidor no e-commerce brasileiro',
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    spacing_after=Pt(18)
)

# ═══════════════════════════════════════════════════════════
# INTRODUÇÃO
# ═══════════════════════════════════════════════════════════
add_heading_bold('Introdução')

add_body(
    'O comércio eletrônico brasileiro vem crescendo de forma consistente nos últimos anos. '
    'Com esse crescimento, multiplicam-se também as situações em que o consumidor tem uma '
    'experiência aquém do esperado — atrasos na entrega, produtos que chegam danificados, '
    'divergência entre o que foi anunciado e o que foi recebido. Essas experiências resultam, '
    'na maioria das vezes, em avaliações negativas nos marketplaces, o que prejudica tanto o '
    'vendedor quanto a plataforma como um todo.'
)

add_body(
    'Do ponto de vista teórico, a satisfação do consumidor pode ser entendida como o resultado '
    'da comparação entre a expectativa prévia e o desempenho percebido do serviço. '
    'Parasuraman et al. (1985) propuseram que lacunas entre esses dois elementos são a principal '
    'causa de percepção de baixa qualidade. No comércio eletrônico, essas lacunas se manifestam '
    'em etapas específicas da jornada de compra — especialmente na logística de entrega, que '
    'frequentemente foge ao controle direto do vendedor. Kotler e Keller (2018) reforçam que a '
    'gestão da experiência do cliente é um dos pilares para a sustentabilidade de qualquer '
    'operação de varejo, e isso vale tanto para o varejo físico quanto para o digital.'
)

add_body(
    'Chevalier e Mayzlin (2006) demonstraram que avaliações online exercem impacto direto sobre '
    'as vendas, o que reforça a importância de se compreender quais fatores levam o consumidor a '
    'atribuir notas baixas. No entanto, a maioria das pesquisas sobre insatisfação em e-commerce '
    'trata o problema de forma binária: o consumidor está ou não está satisfeito. Essa simplificação '
    'ignora a intensidade do fenômeno. Um vendedor com 3 avaliações negativas em 50 transações '
    'enfrenta uma realidade completamente diferente daquele que acumula 25 avaliações negativas '
    'em 100 transações, embora ambos possam ser classificados como "insatisfatórios" em uma '
    'análise binária. A contagem de reviews negativos é uma variável discreta e não-negativa, e '
    'exige tratamento estatístico adequado.'
)

add_body(
    'Para variáveis com essa natureza, os modelos de regressão para dados de contagem são o '
    'ferramental apropriado. O modelo de Poisson é o ponto de partida, mas assume que a variância '
    'condicional é igual à média — pressuposto frequentemente violado em dados reais, configurando '
    'o que se conhece como sobredispersão. Quando isso acontece, a regressão Binomial Negativa '
    'se torna a alternativa natural, pois acomoda essa variabilidade adicional por meio de um '
    'parâmetro extra (Cameron e Trivedi, 2013; Hilbe, 2011). Fávero e Belfiore (2017) detalham a '
    'aplicação prática desses modelos em contextos econômicos e de negócios, incluindo os testes '
    'diagnósticos necessários para a seleção entre Poisson e Binomial Negativa.'
)

add_body(
    'Este trabalho propõe investigar os determinantes da insatisfação do consumidor no e-commerce '
    'brasileiro. A variável de interesse será a contagem de reviews negativos (notas 1 e 2) por '
    'vendedor, e os dados utilizados serão provenientes da base pública da Olist (Olist, 2018), que '
    'contém aproximadamente 100 mil transações realizadas entre 2016 e 2018 em marketplaces '
    'brasileiros. O foco da análise recairá sobre fatores operacionais do vendedor — como atraso na '
    'entrega, valor do frete e categoria de produto — e seu papel como determinantes da taxa de '
    'insatisfação.'
)

# ═══════════════════════════════════════════════════════════
# OBJETIVO
# ═══════════════════════════════════════════════════════════
add_heading_bold('Objetivo')

add_body(
    'Identificar os determinantes da insatisfação do consumidor no e-commerce brasileiro, '
    'operacionalizada como a contagem de reviews negativos (notas 1 e 2) por vendedor, '
    'por meio da aplicação de modelos de regressão para dados de contagem — Poisson e '
    'Binomial Negativa — a dados transacionais públicos da plataforma Olist.'
)

# ═══════════════════════════════════════════════════════════
# METODOLOGIA
# ═══════════════════════════════════════════════════════════
add_heading_bold('Metodologia')

add_body(
    'A pesquisa possui caráter quantitativo e descritivo, apoiada em dados secundários de '
    'acesso público. A base de dados utilizada será o Brazilian E-Commerce Public Dataset, '
    'disponibilizado pela empresa Olist na plataforma Kaggle (Olist, 2018). O dataset contém '
    'registros de aproximadamente 100 mil pedidos realizados em marketplaces brasileiros '
    'entre 2016 e 2018, distribuídos em sete tabelas relacionais: pedidos, itens de pedido, '
    'pagamentos, avaliações, produtos, vendedores e tradução de categorias de produto.'
)

add_body(
    'A unidade de análise será o vendedor. Os dados das sete tabelas serão cruzados e '
    'agregados por identificador do vendedor, gerando uma base cross-section. As variáveis '
    'do modelo serão organizadas da seguinte forma:'
)

# Variáveis em formato de lista
items = [
    ('Variável dependente (Y): ', False,
     'contagem de reviews negativos (score 1 ou 2) por vendedor.'),
    ('Offset: ', False,
     'logaritmo natural do total de reviews recebidos pelo vendedor. '
     'Com o offset, o modelo estimará a taxa de insatisfação, e não a contagem absoluta, '
     'controlando pelo volume de vendas de cada vendedor.'),
    ('Variáveis independentes: ', False,
     'atraso médio na entrega (em dias), ticket médio (R$), frete médio (R$), '
     'peso médio do produto (g), categoria principal do vendedor e unidade federativa '
     'do vendedor.'),
]

for label, _, desc in items:
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run = p.add_run(desc)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)

add_body('')  # small spacer

add_body(
    'Serão incluídos na amostra apenas vendedores com no mínimo 5 reviews, de modo a '
    'garantir representatividade estatística na estimativa da taxa de insatisfação. '
    'As variáveis categóricas (categoria de produto e estado) serão transformadas em '
    'variáveis dummy, agrupando-se categorias com menos de 30 observações em um grupo '
    'residual denominado "other".'
)

add_body(
    'A modelagem seguirá o procedimento descrito por Cameron e Trivedi (2013) e '
    'Fávero e Belfiore (2017). Primeiro, será estimado um modelo de regressão Poisson como '
    'baseline, utilizando o método de mínimos quadrados iterativamente reponderados (IRLS). '
    'Em seguida, será aplicado o teste de sobredispersão de Cameron e Trivedi para verificar '
    'se a variância condicional excede a média condicional. Caso a sobredispersão seja '
    'confirmada, será estimado um modelo de regressão Binomial Negativa (NB2), com o '
    'parâmetro de dispersão (alfa) estimado pelo método dos momentos. A comparação '
    'formal entre os dois modelos será feita pelo teste de razão de verossimilhança.'
)

add_body(
    'A seleção do modelo final considerará os critérios de informação de Akaike (AIC) e '
    'bayesiano (BIC), além da estatística qui-quadrado de Pearson padronizada pelos graus '
    'de liberdade — valores próximos a 1 indicam bom ajuste (Hilbe, 2011). Variáveis com '
    'p-valor superior a 0,05 serão removidas do modelo completo para obtenção de um modelo '
    'reduzido mais parcimonioso.'
)

add_body(
    'A interpretação dos coeficientes será feita por meio de exp(β), que representa o fator '
    'multiplicativo sobre a taxa esperada de reviews negativos para cada incremento unitário '
    'na variável independente. As análises serão conduzidas em Python, com uso das '
    'bibliotecas pandas para manipulação de dados, statsmodels para estimação dos modelos '
    'lineares generalizados e scipy para testes estatísticos.'
)

# ═══════════════════════════════════════════════════════════
# RESULTADOS ESPERADOS
# ═══════════════════════════════════════════════════════════
add_heading_bold('Resultados Esperados')

add_body(
    'Espera-se que o modelo Binomial Negativa apresente ajuste superior ao modelo de '
    'Poisson, dado o comportamento típico de sobredispersão em variáveis de contagem '
    'provenientes de dados transacionais. A partir da estimação e interpretação dos '
    'coeficientes, pretende-se identificar quais fatores operacionais — como atraso na '
    'entrega, valor do frete e categoria de produto — estão estatisticamente associados '
    'a uma maior taxa de insatisfação do consumidor.'
)

add_body(
    'Do ponto de vista prático, os resultados poderão indicar alavancas de atuação para '
    'vendedores e plataformas de e-commerce que busquem reduzir a incidência de '
    'avaliações negativas. Se, por exemplo, o atraso na entrega se confirmar como '
    'determinante relevante, isso reforçaria a necessidade de investimentos em logística '
    'e gestão de prazos como estratégia de melhoria da experiência do consumidor.'
)

# ═══════════════════════════════════════════════════════════
# CRONOGRAMA DE ATIVIDADES
# ═══════════════════════════════════════════════════════════
add_heading_bold('Cronograma de Atividades')

# Tabela do cronograma
atividades = [
    ('Revisão bibliográfica',              ['X', 'X', ' ', ' ', ' ', ' ']),
    ('Coleta e preparação dos dados',      ['X', 'X', ' ', ' ', ' ', ' ']),
    ('Análise exploratória',               [' ', 'X', 'X', ' ', ' ', ' ']),
    ('Modelagem estatística',              [' ', ' ', 'X', 'X', ' ', ' ']),
    ('Entrega: Resultados Preliminares',   [' ', ' ', ' ', 'X', ' ', ' ']),
    ('Redação do TCC',                     [' ', ' ', ' ', 'X', 'X', ' ']),
    ('Revisão e entrega final do TCC',     [' ', ' ', ' ', ' ', 'X', 'X']),
]
meses = ['Mar/26', 'Abr/26', 'Mai/26', 'Jun/26', 'Jul/26', 'Ago/26']

table = doc.add_table(rows=1 + len(atividades), cols=1 + len(meses))
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Estilo da tabela
table.style = 'Table Grid'

# Header
hdr = table.rows[0]
hdr.cells[0].text = 'Atividade'
for j, mes in enumerate(meses):
    hdr.cells[j + 1].text = mes

# Dados
for i, (ativ, marks) in enumerate(atividades):
    row = table.rows[i + 1]
    row.cells[0].text = ativ
    for j, mark in enumerate(marks):
        row.cells[j + 1].text = mark
        # Centralizar o X
        for paragraph in row.cells[j + 1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Formatar fonte da tabela
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)

# Header em negrito
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════
# REFERÊNCIAS
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()  # espaço
add_heading_bold('Referências')

refs = [
    'Cameron, A.C.; Trivedi, P.K. 2013. Regression Analysis of Count Data. 2ed. '
    'Cambridge University Press, Cambridge, Reino Unido.',

    'Chevalier, J.A.; Mayzlin, D. 2006. The effect of word of mouth on sales: '
    'online book reviews. Journal of Marketing Research 43(3): 345-354.',

    'Fávero, L.P.; Belfiore, P. 2017. Manual de Análise de Dados: Estatística e '
    'modelagem multivariada com Excel, SPSS e Stata. Elsevier, Rio de Janeiro, RJ, Brasil.',

    'Hilbe, J.M. 2011. Negative Binomial Regression. 2ed. Cambridge University Press, '
    'Cambridge, Reino Unido.',

    'Kotler, P.; Keller, K.L. 2018. Administração de Marketing. 15ed. Pearson Education '
    'do Brasil, São Paulo, SP, Brasil.',

    'Olist. 2018. Brazilian E-Commerce Public Dataset by Olist. Disponível em: '
    '<https://www.kaggle.com/datasets/olistbr/brazilian-ecommerce>. Acesso em: 20 mar. 2026.',

    'Parasuraman, A.; Zeithaml, V.A.; Berry, L.L. 1985. A conceptual model of service '
    'quality and its implications for future research. Journal of Marketing 49(4): 41-50.',
]

for ref in refs:
    add_ref(ref)

# ── Salvar ──
out_path = 'Projeto_de_Pesquisa.docx'
doc.save(out_path)
print(f'Salvo: {out_path}')
